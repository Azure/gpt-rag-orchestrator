import json
import logging
import os
import re
import html
from datetime import datetime, timedelta, timezone
from azure.storage.blob import (
    BlobServiceClient, 
    ContentSettings, 
    generate_blob_sas,
    BlobSasPermissions
)
from azure.identity import DefaultAzureCredential
from shared.util import get_conversation
import markdown
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import io

def parse_markdown_to_html(text):
    """
    Convert markdown text to HTML with proper formatting.
    Falls back to basic formatting if markdown library is not available.
    
    Args:
        text: Raw text that may contain markdown
        
    Returns:
        str: HTML formatted text
    """
    if not text:
        return ""
    
    # Escape HTML first to prevent XSS
    text = html.escape(text)

    # Use markdown library with extensions for better formatting
    md = markdown.Markdown(extensions=[
        'fenced_code',
        'tables', 
        'toc',
        'nl2br',  # Convert newlines to <br>
        'sane_lists'
    ])
    return md.convert(text)


def format_conversation_as_html(conversation_data):
    """
    Convert conversation data to a readable HTML format.
    
    Args:
        conversation_data: The conversation data from Cosmos DB
        
    Returns:
        str: HTML formatted conversation
    """
    html_template = """
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Shared Conversation</title>
        <style>
            * {{
                box-sizing: border-box;
            }}
            
            body {{
                font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
                max-width: 900px;
                margin: 0 auto;
                padding: 32px;
                background-color: #fafbfc;
                color: #1f2937;
                line-height: 1.6;
                min-height: 100vh;
            }}
            
            .conversation-header {{
                background: white;
                padding: 32px;
                border-radius: 16px;
                margin-bottom: 32px;
                border: 1px solid #e5e7eb;
                box-shadow: 0 1px 3px rgba(0, 0, 0, 0.05);
            }}
            
            .conversation-header h1 {{
                margin: 0 0 20px 0;
                font-size: 28px;
                font-weight: 600;
                color: #111827;
                letter-spacing: -0.025em;
            }}
            
            .conversation-header p {{
                margin: 8px 0;
                color: #6b7280;
                font-size: 15px;
                font-weight: 500;
            }}
            
            .conversation-header strong {{
                color: #374151;
                font-weight: 600;
            }}
            
            .message {{
                background: white;
                margin: 20px 0;
                padding: 24px;
                border-radius: 12px;
                border: 1px solid #e5e7eb;
                box-shadow: 0 1px 3px rgba(0, 0, 0, 0.05);
                transition: all 0.2s cubic-bezier(0.4, 0, 0.2, 1);
            }}
            
            .message:hover {{
                border-color: #d1d5db;
                box-shadow: 0 4px 12px rgba(0, 0, 0, 0.08);
                transform: translateY(-1px);
            }}
            
            .user-message {{
                background: #f8fafc;
                border-left: 3px solid #3b82f6;
                margin-left: 32px;
            }}
            
            .freddaid-message {{
                background: #f9fafb;
                border-left: 3px solid #6b7280;
                margin-right: 32px;
            }}
            
            .role {{
                font-weight: 600;
                font-size: 13px;
                margin-bottom: 12px;
                display: flex;
                align-items: center;
                gap: 8px;
                letter-spacing: 0.025em;
                color: #374151;
                text-transform: uppercase;
            }}
            
            .role::before {{
                content: '';
                width: 6px;
                height: 6px;
                border-radius: 50%;
                background: #9ca3af;
            }}
            
            .user-message .role::before {{
                background: #3b82f6;
            }}
            
            .freddaid-message .role::before {{
                background: #6b7280;
            }}
            
            .content {{
                line-height: 1.7;
                color: #1f2937;
                font-size: 15px;
                font-weight: 400;
            }}
            
            /* Markdown formatting styles */
            .content h1, .content h2, .content h3, .content h4, .content h5, .content h6 {{
                margin: 20px 0 12px 0;
                font-weight: 600;
                color: #111827;
                line-height: 1.3;
            }}
            
            .content h1 {{ font-size: 24px; border-bottom: 2px solid #e5e7eb; padding-bottom: 8px; }}
            .content h2 {{ font-size: 20px; border-bottom: 1px solid #e5e7eb; padding-bottom: 6px; }}
            .content h3 {{ font-size: 18px; }}
            .content h4 {{ font-size: 16px; }}
            .content h5 {{ font-size: 15px; }}
            .content h6 {{ font-size: 14px; }}
            
            .content p {{
                margin: 12px 0;
            }}
            
            .content strong {{
                font-weight: 600;
                color: #111827;
            }}
            
            .content em {{
                font-style: italic;
                color: #374151;
            }}
            
            .content code {{
                background-color: #f3f4f6;
                color: #dc2626;
                padding: 2px 6px;
                border-radius: 4px;
                font-family: 'Monaco', 'Menlo', 'Ubuntu Mono', monospace;
                font-size: 13px;
                border: 1px solid #e5e7eb;
            }}
            
            .content pre {{
                background-color: #1f2937;
                color: #f9fafb;
                padding: 16px;
                border-radius: 8px;
                overflow-x: auto;
                margin: 16px 0;
                border: 1px solid #374151;
            }}
            
            .content pre code {{
                background: none;
                color: inherit;
                padding: 0;
                border: none;
                font-size: 14px;
            }}
            
            .content ul, .content ol {{
                margin: 12px 0;
                padding-left: 24px;
            }}
            
            .content li {{
                margin: 6px 0;
                line-height: 1.6;
            }}
            
            .content ul li {{
                list-style-type: disc;
            }}
            
            .content ol li {{
                list-style-type: decimal;
            }}
            
            .content blockquote {{
                border-left: 4px solid #d1d5db;
                padding-left: 16px;
                margin: 16px 0;
                font-style: italic;
                color: #6b7280;
                background-color: #f9fafb;
                padding: 12px 16px;
                border-radius: 4px;
            }}
            
            .content a {{
                color: #3b82f6;
                text-decoration: none;
                border-bottom: 1px solid transparent;
                transition: border-color 0.2s;
            }}
            
            .content a:hover {{
                border-bottom-color: #3b82f6;
            }}
            
            .content table {{
                border-collapse: collapse;
                width: 100%;
                margin: 16px 0;
                border: 1px solid #e5e7eb;
                border-radius: 6px;
                overflow: hidden;
            }}
            
            .content th, .content td {{
                border: 1px solid #e5e7eb;
                padding: 8px 12px;
                text-align: left;
            }}
            
            .content th {{
                background-color: #f9fafb;
                font-weight: 600;
                color: #374151;
            }}
            
            .content tr:nth-child(even) {{
                background-color: #f9fafb;
            }}
            
            .content hr {{
                border: none;
                height: 1px;
                background-color: #e5e7eb;
                margin: 24px 0;
            }}
            
            .timestamp {{
                color: #9ca3af;
                font-size: 12px;
                margin-top: 16px;
                font-weight: 500;
            }}
            
            .export-info {{
                background: #f3f4f6;
                padding: 16px 20px;
                border-radius: 12px;
                margin-bottom: 32px;
                font-size: 14px;
                color: #4b5563;
                border: 1px solid #e5e7eb;
                display: flex;
                align-items: center;
                gap: 8px;
                font-weight: 500;
            }}
            
            .messages {{
                animation: fadeIn 0.6s ease-out;
            }}
            
            @keyframes fadeIn {{
                from {{ opacity: 0; transform: translateY(20px); }}
                to {{ opacity: 1; transform: translateY(0); }}
            }}
            
            /* Clean scrollbar */
            ::-webkit-scrollbar {{
                width: 6px;
            }}
            
            ::-webkit-scrollbar-track {{
                background: #f1f5f9;
            }}
            
            ::-webkit-scrollbar-thumb {{
                background: #cbd5e1;
                border-radius: 3px;
            }}
            
            ::-webkit-scrollbar-thumb:hover {{
                background: #94a3b8;
            }}
            
            @media (max-width: 768px) {{
                body {{
                    padding: 20px;
                }}
                
                .conversation-header {{
                    padding: 24px;
                }}
                
                .message {{
                    padding: 20px;
                }}
                
                .user-message {{
                    margin-left: 16px;
                }}
                
                .freddaid-message {{
                    margin-right: 16px;
                }}
            }}
            
            @media (max-width: 480px) {{
                .user-message {{
                    margin-left: 8px;
                }}
                
                .freddaid-message {{
                    margin-right: 8px;
                }}
            }}
        </style>
    </head>
    <body>
        <div class="export-info">
            📋 This conversation was exported on {export_date}
        </div>
        
        <div class="conversation-header">
            <h1>Conversation Export</h1>
            <p><strong>Started:</strong> {start_date}</p>
            <p><strong>Conversation ID:</strong> {conversation_id}</p>
            <p><strong>Total Messages:</strong> {message_count}</p>
        </div>
        
        <div class="messages">
            {messages_html}
        </div>
    </body>
    </html>
    """
    
    messages_html = ""
    messages = conversation_data.get('messages', [])
    for message in messages:
        role = message.get('role', 'unknown')
        content = message.get('content', '')
        
        # Parse markdown content to HTML
        formatted_content = parse_markdown_to_html(content)
        
        css_class = 'user-message' if role == 'user' else 'freddaid-message'
        role_display = 'User' if role == 'user' else 'Freddaid'
        
        messages_html += f"""
        <div class="message {css_class}">
            <div class="role">{role_display}</div>
            <div class="content">{formatted_content}</div>
        </div>
        """
    
    return html_template.format(
        export_date=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        start_date=conversation_data.get('start_date', 'Unknown'),
        conversation_id=conversation_data.get('id', 'Unknown'),
        message_count=len(messages),
        messages_html=messages_html
    )

def format_conversation_as_json(conversation_data):
    """
    Convert conversation data to formatted JSON.
    
    Args:
        conversation_data: The conversation data from Cosmos DB
        
    Returns:
        str: JSON formatted conversation
    """
    messages = conversation_data.get('messages', [])
    export_data = {
        "conversation_id": conversation_data.get('id'),
        "export_date": datetime.now().isoformat(),
        "start_date": conversation_data.get('start_date'),
        "message_count": len(messages),
        "messages": []
    }
    
    # Add messages without sensitive user data
    for message in messages:
        export_message = {
            "role": message.get('role'),
            "content": message.get('content')
        }
        export_data["messages"].append(export_message)
    
    return json.dumps(export_data, indent=2, ensure_ascii=False)

def parse_markdown_to_docx_content(text, paragraph):
    """
    Parse simple markdown content and apply basic formatting to a Word paragraph.
    This handles basic markdown like **bold**, *italic*, and `code`.
    
    Args:
        text: Text that may contain markdown
        paragraph: Word document paragraph object to add content to
    """
    if not text:
        return
    
    # Simple regex patterns for basic markdown
    # This is a simplified approach - for complex markdown, you'd want a proper parser
    
    # Split by markdown patterns while preserving the patterns
    
    # Pattern to match **bold**, *italic*, `code`, and plain text
    pattern = r'(\*\*.*?\*\*|\*.*?\*|`.*?`)'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
            
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            # Italic text
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            # Code text
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            # Plain text
            paragraph.add_run(part)

def process_message_content(doc, content, role, user_style, assistant_style):
    """
    Process message content and add it to the document with proper formatting.
    Handles headers, regular text, and inline formatting.
    
    Args:
        doc: Word document object
        content: Message content text
        role: Message role ('user' or 'assistant')
        user_style: Style for user messages
        assistant_style: Style for assistant messages
    """
    if not content:
        return
    
    
    # Split content by lines to handle multi-line messages
    content_lines = content.split('\n')
    
    for line_idx, line in enumerate(content_lines):
        line = line.strip()
        if not line:
            # Add empty paragraph for spacing
            doc.add_paragraph()
            continue
        
        # Check if line is a header (starts with #)
        header_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if header_match:
            # This is a header
            header_level = len(header_match.group(1))  # Number of # symbols
            header_text = header_match.group(2)
            
            # Add header to document (level 1-6 maps to Word heading levels 1-6)
            header_paragraph = doc.add_heading(header_text, level=min(header_level, 6))
            header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            # Regular paragraph
            p = doc.add_paragraph()
            parse_markdown_to_docx_content(line, p)
            try:
                if role == 'user':
                    p.style = user_style
                else:
                    p.style = assistant_style
            except:
                pass  # Use default style if custom style fails

def format_conversation_as_docx(conversation_data):
    """
    Convert conversation data to a Word document (.docx).
    
    Args:
        conversation_data: The conversation data from Cosmos DB
        
    Returns:
        bytes: Word document as bytes
    """
    
    # Create a new Word document
    doc = Document()
    
    # Set up document styles
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    
    # Create custom styles for user and assistant messages
    try:
        user_style = doc.styles.add_style('UserMessage', WD_STYLE_TYPE.PARAGRAPH)
        user_style.base_style = doc.styles['Normal']
        user_style.font.name = 'Calibri'
        user_style.font.size = Pt(11)
        user_style.paragraph_format.left_indent = Inches(0.5)
        user_style.paragraph_format.space_after = Pt(6)
    except:
        user_style = doc.styles['Normal']  # Fallback if style exists
    
    try:
        assistant_style = doc.styles.add_style('AssistantMessage', WD_STYLE_TYPE.PARAGRAPH)
        assistant_style.base_style = doc.styles['Normal']
        assistant_style.font.name = 'Calibri'
        assistant_style.font.size = Pt(11)
        assistant_style.paragraph_format.right_indent = Inches(0.5)
        assistant_style.paragraph_format.space_after = Pt(6)
    except:
        assistant_style = doc.styles['Normal']  # Fallback if style exists
    
    # Add document header
    header = doc.add_heading('Conversation Export', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add conversation metadata
    doc.add_paragraph(f"Export Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Conversation ID: {conversation_data.get('id', 'Unknown')}")
    doc.add_paragraph(f"Started: {conversation_data.get('start_date', 'Unknown')}")
    doc.add_paragraph(f"Total Messages: {len(conversation_data.get('messages', []))}")
    
    # Add a separator
    doc.add_paragraph("=" * 50)
    doc.add_paragraph()
    
    # Add messages
    messages = conversation_data.get('messages', [])
    for i, message in enumerate(messages):
        role = message.get('role', 'unknown')
        content = message.get('content', '')
        
        role_display = 'User' if role == 'user' else 'Freddaid'
        
        # Check if content starts with a header or if it's a short single-line message
        content_lines = content.split('\n')
        first_line = content_lines[0].strip() if content_lines else ''
        is_first_line_header = re.match(r'^#{1,6}\s+', first_line)
        
        if role == 'user' and not is_first_line_header and len(content_lines) == 1 and len(first_line) < 100:
            # For short user messages, put on same line as "User:"
            role_paragraph = doc.add_paragraph()
            role_run = role_paragraph.add_run(f"{role_display}: ")
            role_run.bold = True
            role_run.font.size = Pt(12)
            
            # Add the content on the same line
            parse_markdown_to_docx_content(first_line, role_paragraph)
            
            # Don't apply user_style for inline messages to avoid indentation
            # Instead, apply basic formatting without indentation
            role_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            role_paragraph.paragraph_format.left_indent = Inches(0)
            role_paragraph.paragraph_format.space_after = Pt(6)
        else:
            # For longer messages or messages with headers, use separate lines
            role_paragraph = doc.add_heading(f"{role_display}:", level=2)
            role_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Process the content with proper header handling
            process_message_content(doc, content, role, user_style, assistant_style)
        
        # Add spacing between messages
        if i < len(messages) - 1:  # Don't add extra space after last message
            doc.add_paragraph()
            doc.add_paragraph("-" * 30)
            doc.add_paragraph()
    
    # Save document to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    return doc_bytes.getvalue()

def upload_to_blob_storage(content, filename, user_id, content_type="text/html"):
    """
    Upload content to Azure Blob Storage and return shareable URL.
    
    Args:
        content: File content to upload
        filename: Name of the file
        content_type: MIME type of the content
        user_id: path to the user's folder in the blob storage
    Returns:
        str: Shareable URL to the uploaded file
    """
    try:
        # Get Azure Storage connection
        storage_account_url = os.environ.get("AZURE_STORAGE_ACCOUNT_URL")
        if not storage_account_url:
            raise ValueError("AZURE_STORAGE_ACCOUNT_URL environment variable not set")
        
        # Use managed identity for authentication
        credential = DefaultAzureCredential()
        blob_service_client = BlobServiceClient(
            account_url=storage_account_url, 
            credential=credential
        )
        
        # Container for shared conversations
        container_name = "shared-conversations"
        
        # Create container if it doesn't exist
        try:
            blob_service_client.create_container(container_name)
        except Exception:
            pass  # Container might already exist
        
        # Upload the file
        blob_client = blob_service_client.get_blob_client(
            container=container_name, 
            blob=f"{user_id}/{filename}"
        )
        
        blob_client.upload_blob(
            content, 
            overwrite=True,
            content_settings=ContentSettings(content_type=content_type)
        )
        
        # Get a user delegation key to sign the SAS token, valid for 7 days
        delegation_key_start_time = datetime.now(timezone.utc)
        delegation_key_expiry_time = delegation_key_start_time + timedelta(days=3650)
        
        user_delegation_key = blob_service_client.get_user_delegation_key(
            key_start_time=delegation_key_start_time,
            key_expiry_time=delegation_key_expiry_time
        )
        
        # Generate a user-delegation SAS token for the blob
        sas_token = generate_blob_sas(
            account_name=blob_service_client.account_name,
            container_name=container_name,
            blob_name=f"{user_id}/{filename}",
            user_delegation_key=user_delegation_key,
            permission=BlobSasPermissions(read=True),
            expiry=delegation_key_expiry_time
        )
        
        # Construct the full URL with SAS token
        blob_url_with_sas = f"{blob_client.url}?{sas_token}"
        
        return blob_url_with_sas
        
    except Exception as e:
        logging.error(f"Error uploading to blob storage: {str(e)}")
        raise

def export_conversation(conversation_id, user_id, export_format="html"):
    """
    Export a conversation to a file and upload to blob storage.
    
    Args:
        conversation_id: ID of the conversation to export
        user_id: ID of the user requesting the export (for security)
        export_format: Format to export ("html", "json")
        
    Returns:
        dict: Export result with URL and metadata
    """
    try:
        # Get conversation data
        conversation_data = get_conversation(conversation_id, user_id)
        
        if not conversation_data:
            raise ValueError("Conversation not found or access denied")
        
        # Add the conversation ID to the data for formatting
        conversation_data['id'] = conversation_id
        
        # Generate filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{conversation_id}/Freddaid_{timestamp}.{export_format}"

        # Format content based on requested format
        if export_format.lower() == "html":
            content = format_conversation_as_html(conversation_data)
            content_type = "text/html"
        elif export_format.lower() == "json":
            content = format_conversation_as_json(conversation_data)
            content_type = "application/json"
        elif export_format.lower() == "docx":
            content = format_conversation_as_docx(conversation_data)
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            raise ValueError(f"Unsupported export format: {export_format}")
        
        # Upload to blob storage
        share_url = upload_to_blob_storage(content, filename, user_id, content_type)
        
        return {
            "success": True,
            "share_url": share_url,
            "filename": filename,
            "format": export_format,
            "message_count": len(conversation_data.get('messages', [])),
            "export_date": datetime.now().isoformat()
        }
        
    except Exception as e:
        logging.error(f"Error exporting conversation {conversation_id}: {str(e)}")
        return {
            "success": False,
            "error": str(e)
        } 